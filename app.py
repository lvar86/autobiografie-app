"""
Autobiografie Assistent — Streamlit webinterface
"""

import os
import json
from datetime import datetime
from pathlib import Path

import streamlit as st
import anthropic
from dotenv import load_dotenv

load_dotenv()

# ---------------------------------------------------------------------------
# Constanten
# ---------------------------------------------------------------------------

MODEL = "claude-opus-4-7"
OUTPUT_DIR = Path("autobiografieën")
CHUNK_SIZE = 6  # berichten per chunk (3 vragen + 3 antwoorden)

INTERVIEW_SYSTEM_PROMPT = """Je bent een warme, empathische autobiografie-interviewer die mensen helpt hun levensverhaal te ontdekken en te vertellen. Je bent tegelijkertijd journalist, therapeut en schrijver — nieuwsgierig, geduldig en creatief.

**Jouw doel:** door slimme vragen de mooiste, meest menselijke en meest betekenisvolle momenten uit iemands leven naar boven halen — momenten die de persoon misschien zelf al vergeten was of nooit zo had verwoord.

**Aanpak:**
- Begin met brede, open vragen om een beeld te krijgen van de persoon
- Verdiep je ALTIJD in interessante details die worden genoemd — vraag door
- Vraag naar emoties, gedachten, geuren, kleuren, geluiden — maak herinneringen levendig
- Ontdek de "waarom" achter keuzes en levensgebeurtenissen
- Verbind thema's: "Je noemde eerder X — hoe past dat bij wat je nu vertelt?"
- Stel vragen die mensen aan het denken zetten: "Wat zou jij je 20-jarige zelf willen zeggen?"

**Regels:**
- Stel ALTIJD slechts ÉÉN vraag per bericht
- Reageer eerst warm en samengevat op wat de persoon deelde, dan pas de nieuwe vraag
- Ga niet te snel — laat de persoon uitweiden
- Bij emotionele onderwerpen: erken eerst, vraag daarna voorzichtig door
- Gebruik de naam van de persoon als je die kent

**Thema's om te verkennen** (niet allemaal tegelijk, maar geleidelijk):
- Jeugd en herkomst: ouders, woonplaats, herinneringen, school
- Bepalende momenten: keuzes die alles veranderden
- Relaties: vriendschappen, liefde, verlies
- Passies en talenten: wat maakt deze persoon uniek?
- Uitdagingen: hoe zijn ze overwonnen?
- Dromen: wat heeft deze persoon nagestreefd?
- Erfenis: wat wil deze persoon achterlaten?

Begin met een hartelijke begroeting en een eerste, open vraag."""


CHAPTER_SYSTEM_PROMPT = """Je bent een professionele biografieschrijver. Je krijgt een fragment van een autobiografie-interview en schrijft er één hoofdstuk van een autobiografie van.

Richtlijnen:
- Schrijf in de eerste persoon ("Ik...")
- Warme, literaire toon — geen droge opsomming
- Verwerk concrete details, anekdotes en emoties uit het fragment
- Maak het levendig: laat de lezer de momenten beleven
- Geef het hoofdstuk een passende, poëtische titel
- Lengte: 300–600 woorden

Formaat — geef ALLEEN dit terug, niets anders:
## [Hoofdstuktitel]

[Hoofdstukinhoud]"""


ASSEMBLY_SYSTEM_PROMPT = """Je bent een professionele redacteur die losse hoofdstukken samenvoegt tot één vloeiende autobiografie.

Jouw taak:
- Verbind de hoofdstukken tot een samenhangend geheel
- Voeg een korte inleiding toe (max 100 woorden)
- Voeg een afsluitende passage toe over erfenis of toekomstblik
- Zorg voor vloeiende overgangen tussen hoofdstukken
- Verander de inhoud van de hoofdstukken zo min mogelijk — behoud de stem
- Schrijf in de eerste persoon

Geef de volledige autobiografie terug, inclusief alle hoofdstukken."""


# ---------------------------------------------------------------------------
# Hulpfuncties
# ---------------------------------------------------------------------------

def get_client():
    api_key = os.environ.get("ANTHROPIC_API_KEY") or st.session_state.get("api_key", "")
    if not api_key:
        return None
    return anthropic.Anthropic(api_key=api_key)


def get_real_messages(messages):
    return [m for m in messages if m["content"] != "Laten we beginnen."]


def get_chunks(messages):
    real = get_real_messages(messages)
    chunks = []
    for i in range(0, len(real), CHUNK_SIZE):
        chunks.append(real[i:i + CHUNK_SIZE])
    return chunks


def chunk_is_complete(chunk_index, messages):
    chunks = get_chunks(messages)
    return chunk_index < len(chunks) - 1


def dropbox_list_sessions(token):
    """Haal lijst van opgeslagen sessies op uit Dropbox."""
    try:
        import dropbox
        dbx = dropbox.Dropbox(token)
        result = dbx.files_list_folder("/Autobiografie")
        sessions = [
            e.name for e in result.entries
            if e.name.startswith("data_") and e.name.endswith(".json")
        ]
        return sorted(sessions, reverse=True)
    except Exception:
        return []


def dropbox_download_session(token, filename):
    """Download en parseer een sessie-JSON uit Dropbox."""
    try:
        import dropbox
        dbx = dropbox.Dropbox(token)
        _, response = dbx.files_download(f"/Autobiografie/{filename}")
        return json.loads(response.content.decode("utf-8"))
    except Exception as e:
        return None, str(e)


def dropbox_connect(token):
    """Verbind met Dropbox en geef accountnaam terug."""
    try:
        import dropbox
        dbx = dropbox.Dropbox(token)
        account = dbx.users_get_current_account()
        return True, account.name.display_name
    except Exception as e:
        return False, str(e)


def dropbox_upload(token, filename, content):
    """Upload een bestand naar Dropbox."""
    try:
        import dropbox
        from dropbox.files import WriteMode
        dbx = dropbox.Dropbox(token)
        path = f"/Autobiografie/{filename}"
        dbx.files_upload(content.encode("utf-8"), path, mode=WriteMode.overwrite)
        return True, path
    except Exception as e:
        return False, str(e)


def cloud_save(filename, content):
    """Sla op in Dropbox als token aanwezig is, toon toast-melding."""
    token = st.session_state.get("dropbox_token", "")
    if not token or not st.session_state.get("dropbox_confirmed"):
        return
    ok, result = dropbox_upload(token, filename, content)
    if ok:
        st.toast(f"☁️ Opgeslagen in Dropbox", icon="✅")
    else:
        st.toast(f"Dropbox fout: {result}", icon="⚠️")


def chapter_to_docx_bytes(chapter_text, chapter_num):
    """Zet hoofdstuktekst om naar Word-bestand als bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # Stijl instellen
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)

    lines = chapter_text.strip().split("\n")
    title = f"Hoofdstuk {chapter_num}"
    body_lines = []

    for line in lines:
        if line.startswith("## "):
            title = line[3:].strip()
        else:
            body_lines.append(line)

    # Titel
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph()

    # Inhoud
    for line in body_lines:
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_chapter_to_dropbox(chapter_text, chapter_num, base_name):
    """Upload een hoofdstuk als Word-bestand naar Dropbox."""
    token = st.session_state.get("dropbox_token", "")
    if not token or not st.session_state.get("dropbox_confirmed"):
        return
    try:
        import dropbox
        from dropbox.files import WriteMode
        dbx = dropbox.Dropbox(token)
        docx_bytes = chapter_to_docx_bytes(chapter_text, chapter_num)
        filename = f"hoofdstuk_{chapter_num:02d}_{base_name}.docx"
        path = f"/Autobiografie/{filename}"
        dbx.files_upload(docx_bytes, path, mode=WriteMode.overwrite)
        st.toast(f"📄 Hoofdstuk {chapter_num} opgeslagen als Word", icon="✅")
    except Exception as e:
        st.toast(f"Word-opslag mislukt: {e}", icon="⚠️")


def autosave(messages, chapters):
    base_name = st.session_state.get("base_name", datetime.now().strftime("%Y%m%d_%H%M%S"))
    data = json.dumps({
        "created_at": datetime.now().isoformat(),
        "messages": messages,
        "chapters": chapters,
    }, ensure_ascii=False, indent=2)

    # Lokaal opslaan (werkt niet op Streamlit Cloud, maar geen fout)
    try:
        OUTPUT_DIR.mkdir(exist_ok=True)
        (OUTPUT_DIR / f"data_{base_name}.json").write_text(data, encoding="utf-8")
    except Exception:
        pass

    # Cloud opslaan
    cloud_save(f"data_{base_name}.json", data)


def save_autobiography(text, base_name):
    content = f"MIJN AUTOBIOGRAFIE\nGegenereerd op: {datetime.now().strftime('%d %B %Y')}\n{'=' * 60}\n\n{text}"

    try:
        OUTPUT_DIR.mkdir(exist_ok=True)
        path = OUTPUT_DIR / f"autobiografie_{base_name}.txt"
        path.write_text(content, encoding="utf-8")
    except Exception:
        path = Path(f"autobiografie_{base_name}.txt")

    cloud_save(f"autobiografie_{base_name}.txt", content)
    return path


def generate_chapter(client, chunk_messages, chunk_index):
    fragment = ""
    for msg in chunk_messages:
        role = "Interviewer" if msg["role"] == "assistant" else "Persoon"
        fragment += f"{role}: {msg['content']}\n\n"

    with client.messages.stream(
        model=MODEL,
        max_tokens=2048,
        thinking={"type": "adaptive"},
        system=[{"type": "text", "text": CHAPTER_SYSTEM_PROMPT,
                 "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user",
                   "content": f"Schrijf een hoofdstuk op basis van dit interviewfragment:\n\n{fragment}"}],
    ) as stream:
        return stream.get_final_text()


def assemble_autobiography(client, chapters):
    all_chapters = "\n\n".join(
        f"Hoofdstuk {i+1}:\n{ch}" for i, ch in enumerate(chapters)
    )
    with client.messages.stream(
        model=MODEL,
        max_tokens=8192,
        thinking={"type": "adaptive"},
        system=[{"type": "text", "text": ASSEMBLY_SYSTEM_PROMPT,
                 "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user",
                   "content": f"Voeg deze hoofdstukken samen tot één autobiografie:\n\n{all_chapters}"}],
    ) as stream:
        return stream.get_final_text()


# ---------------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------------

st.set_page_config(
    page_title="Autobiografie Assistent",
    page_icon="📖",
    layout="wide",
)

st.markdown("""
<style>
    .block-container { max-width: 1100px; }
    h1 { font-size: 1.8rem !important; }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Sessiestatus initialiseren
# ---------------------------------------------------------------------------

defaults = {
    "messages": [],
    "chapters": {},
    "autobiography": None,
    "base_name": datetime.now().strftime("%Y%m%d_%H%M%S"),
    "started": False,
    "api_key": os.environ.get("ANTHROPIC_API_KEY", ""),
    "dropbox_token": "",
    "dropbox_confirmed": False,
    "dropbox_account_name": "",
}
for key, val in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ---------------------------------------------------------------------------
# API-sleutel controleren
# ---------------------------------------------------------------------------

if not st.session_state.api_key:
    st.error("❌ Geen API-sleutel gevonden. Voeg `ANTHROPIC_API_KEY=sk-ant-...` toe aan het `.env` bestand.")
    st.stop()

client = get_client()

# ---------------------------------------------------------------------------
# Zijbalk
# ---------------------------------------------------------------------------

with st.sidebar:
    st.header("☁️ Cloud opslag")

    if st.session_state.dropbox_confirmed:
        st.success(f"✅ Verbonden als **{st.session_state.dropbox_account_name}**")
        st.caption("Alles wordt automatisch opgeslagen in `/Autobiografie/` in jouw Dropbox.")

        if st.button("🧪 Test opslag", use_container_width=True):
            with st.spinner("Testbestand uploaden..."):
                ok, result = dropbox_upload(
                    st.session_state.dropbox_token,
                    "test.txt",
                    f"Verbindingstest — {datetime.now().isoformat()}"
                )
            if ok:
                st.success(f"✅ Gelukt! Bestand staat op {result}")
            else:
                st.error(f"❌ Mislukt: {result}")

        st.divider()
        st.markdown("**📂 Sessie laden**")
        sessions = dropbox_list_sessions(st.session_state.dropbox_token)
        if sessions:
            def format_session(name):
                # data_20260513_143000.json → 13-05-2026 14:30
                try:
                    ts = name.replace("data_", "").replace(".json", "")
                    dt = datetime.strptime(ts, "%Y%m%d_%H%M%S")
                    return dt.strftime("%d-%m-%Y %H:%M")
                except Exception:
                    return name

            selected = st.selectbox(
                "Kies een eerdere sessie",
                options=sessions,
                format_func=format_session,
                index=0,
            )
            if st.button("📂 Laden en verder gaan", use_container_width=True):
                with st.spinner("Sessie laden..."):
                    data = dropbox_download_session(st.session_state.dropbox_token, selected)
                if data:
                    base = selected.replace("data_", "").replace(".json", "")
                    st.session_state.messages = data.get("messages", [])
                    st.session_state.chapters = {
                        int(k): v for k, v in data.get("chapters", {}).items()
                    }
                    st.session_state.autobiography = data.get("autobiography")
                    st.session_state.base_name = base
                    st.session_state.started = True
                    st.success("✅ Sessie geladen!")
                    st.rerun()
                else:
                    st.error("Kon sessie niet laden.")
        else:
            st.caption("Nog geen opgeslagen sessies gevonden.")

        st.divider()
        if st.button("🔌 Ontkoppelen", use_container_width=True):
            st.session_state.dropbox_token = ""
            st.session_state.dropbox_confirmed = False
            st.session_state.dropbox_account_name = ""
            st.rerun()
    else:
        st.caption("Verbind je Dropbox om alles automatisch op te slaan.")
        token_input = st.text_input(
            "Toegangstoken",
            type="password",
            placeholder="sl.xxxxxxxx...",
        )
        if st.button("Verbinden", use_container_width=True, type="primary"):
            if token_input:
                with st.spinner("Verbinden met Dropbox..."):
                    ok, result = dropbox_connect(token_input)
                if ok:
                    st.session_state.dropbox_token = token_input
                    st.session_state.dropbox_confirmed = True
                    st.session_state.dropbox_account_name = result
                    st.rerun()
                else:
                    st.error(f"Verbinding mislukt: {result}")
            else:
                st.warning("Voer eerst een token in.")

        with st.expander("Hoe maak ik een token aan?"):
            st.markdown("""
1. Ga naar [dropbox.com/developers](https://www.dropbox.com/developers/apps)
2. Klik **Create app**
3. Kies **Scoped access** → **Full Dropbox**
4. Geef een naam → **Create app**
5. Tabblad **Settings** → scroll naar **Generated access token** → klik **Generate**
6. Plak de token hierboven en klik Verbinden
""")

    st.divider()
    st.caption(f"Sessie: {st.session_state.base_name}")
    st.caption(f"Berichten: {max(0, len(st.session_state.messages) - 1)}")

# ---------------------------------------------------------------------------
# Layout: twee kolommen
# ---------------------------------------------------------------------------

col_chat, col_book = st.columns([3, 2], gap="large")

# ---------------------------------------------------------------------------
# Linker kolom: het gesprek
# ---------------------------------------------------------------------------

with col_chat:
    st.title("📖 Autobiografie Assistent")
    st.caption("Claude stelt je vragen om jouw levensverhaal te ontdekken.")

    if not st.session_state.started:
        st.session_state.started = True
        with st.spinner("Claude bereidt zich voor..."):
            interview_system = [{"type": "text", "text": INTERVIEW_SYSTEM_PROMPT,
                                 "cache_control": {"type": "ephemeral"}}]
            with client.messages.stream(
                model=MODEL,
                max_tokens=1024,
                thinking={"type": "adaptive"},
                system=interview_system,
                messages=[{"role": "user", "content": "Laten we beginnen."}],
            ) as stream:
                opening = stream.get_final_text()
        st.session_state.messages.append({"role": "user", "content": "Laten we beginnen."})
        st.session_state.messages.append({"role": "assistant", "content": opening})
        autosave(st.session_state.messages, st.session_state.chapters)

    for msg in st.session_state.messages:
        if msg["content"] == "Laten we beginnen.":
            continue
        role = "assistant" if msg["role"] == "assistant" else "user"
        with st.chat_message(role, avatar="📖" if role == "assistant" else "👤"):
            st.markdown(msg["content"])

    if prompt := st.chat_input("Jouw antwoord..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user", avatar="👤"):
            st.markdown(prompt)

        interview_system = [{"type": "text", "text": INTERVIEW_SYSTEM_PROMPT,
                             "cache_control": {"type": "ephemeral"}}]
        with st.chat_message("assistant", avatar="📖"):
            placeholder = st.empty()
            full_response = ""
            with client.messages.stream(
                model=MODEL,
                max_tokens=1024,
                thinking={"type": "adaptive"},
                system=interview_system,
                messages=st.session_state.messages,
            ) as stream:
                for text in stream.text_stream:
                    full_response += text
                    placeholder.markdown(full_response + "▌")
            placeholder.markdown(full_response)

        st.session_state.messages.append({"role": "assistant", "content": full_response})
        autosave(st.session_state.messages, st.session_state.chapters)
        st.rerun()

# ---------------------------------------------------------------------------
# Rechter kolom: het boek
# ---------------------------------------------------------------------------

with col_book:
    st.subheader("📚 Jouw boek")

    chunks = get_chunks(st.session_state.messages)

    if not chunks:
        st.caption("Nog geen gesprek om om te zetten.")
    else:
        for i, chunk in enumerate(chunks):
            is_complete = chunk_is_complete(i, st.session_state.messages)
            chapter_exists = i in st.session_state.chapters
            q_count = sum(1 for m in chunk if m["role"] == "user")

            label = f"Fragment {i + 1} · {q_count} vragen"
            status = "✅" if chapter_exists else ("🔒" if not is_complete else "⬜")

            with st.expander(f"{status} {label}", expanded=chapter_exists):
                if chapter_exists:
                    st.markdown(st.session_state.chapters[i])
                    if st.button("↩️ Opnieuw genereren", key=f"regen_{i}"):
                        with st.spinner(f"Hoofdstuk {i + 1} herschrijven..."):
                            st.session_state.chapters[i] = generate_chapter(client, chunk, i)
                        autosave(st.session_state.messages, st.session_state.chapters)
                        save_chapter_to_dropbox(st.session_state.chapters[i], i + 1, st.session_state.base_name)
                        st.rerun()
                elif is_complete:
                    st.caption("Fragment afgerond — klaar om te schrijven.")
                    if st.button(f"✍️ Schrijf hoofdstuk {i + 1}", key=f"gen_{i}"):
                        with st.spinner(f"Hoofdstuk {i + 1} schrijven..."):
                            st.session_state.chapters[i] = generate_chapter(client, chunk, i)
                        autosave(st.session_state.messages, st.session_state.chapters)
                        save_chapter_to_dropbox(st.session_state.chapters[i], i + 1, st.session_state.base_name)
                        st.rerun()
                else:
                    st.caption(f"Gesprek loopt nog — komt beschikbaar na {CHUNK_SIZE - len(chunk)} berichten.")

    st.divider()

    complete_chapters = [st.session_state.chapters[i] for i in sorted(st.session_state.chapters)]
    can_assemble = len(complete_chapters) >= 1

    if st.button("📖 Samenvoegen tot autobiografie", use_container_width=True,
                 disabled=not can_assemble):
        with st.spinner("Autobiografie samenvoegen..."):
            st.session_state.autobiography = assemble_autobiography(client, complete_chapters)
        save_autobiography(st.session_state.autobiography, st.session_state.base_name)
        autosave(st.session_state.messages, st.session_state.chapters)
        st.rerun()

    if st.session_state.autobiography:
        st.markdown("### 📜 Autobiografie")
        st.markdown(st.session_state.autobiography)
        st.download_button(
            "⬇️ Download als tekstbestand",
            data=st.session_state.autobiography,
            file_name=f"autobiografie_{st.session_state.base_name}.txt",
            mime="text/plain",
            use_container_width=True,
        )
